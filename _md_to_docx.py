"""Convert ספר_פרויקט_Driver_Moodle_FINAL_v2.md → polished docx with Hebrew RTL.

Output styling:
- A4 page, 2cm margins
- Cover-style H1 in cyan/large
- H2 with bottom border in dark blue
- H3 in orange
- Body Arial 11
- Code blocks: Consolas 9 with light grey shaded background + grey border
- Tables: Light Grid Accent 1, header row bolded
- Images: 6.5" width, centered, italic caption
- Blockquotes: italic + grey
"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"D:\yudb"
# Process both books: Hebrew master + English appendix.
JOBS = [
    (os.path.join(ROOT, "ספר_פרויקט_B_Managed_HE.md"),
     os.path.join(ROOT, "ספר_פרויקט_B_Managed_HE.docx")),
    (os.path.join(ROOT, "BManaged_ProjectBook_EN.md"),
     os.path.join(ROOT, "BManaged_ProjectBook_EN.docx")),
]
MD_PATH, DOCX_PATH = JOBS[0]  # default; overwritten in build()

CYAN   = RGBColor(0x00, 0x9F, 0xCC)
ORANGE = RGBColor(0xC0, 0x70, 0x10)
DARK   = RGBColor(0x1F, 0x3D, 0x6B)


def add_bidi(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))


def set_cell_rtl(cell):
    for p in cell.paragraphs:
        add_bidi(p)


def add_heading(doc, text, level):
    if level == 0:
        # cover-style H1
        p = doc.add_paragraph()
        add_bidi(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = CYAN
        return p
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    p = doc.add_paragraph(style=style_map.get(level, "Heading 3"))
    add_bidi(p)
    run = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs or p.runs[0].text != text:
        p.runs[0].text = text
    color = {1: DARK, 2: CYAN, 3: ORANGE}.get(level, DARK)
    size  = {1: 22, 2: 16, 3: 13}.get(level, 12)
    for r in p.runs:
        r.font.color.rgb = color
        r.font.size = Pt(size)
        r.bold = True
    if level == 1:
        # bottom border on H1 mid-doc
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "00B7E0")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    add_bidi(p)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.4)
    pf.right_indent = Cm(0.4)
    pf.space_before = Pt(4)
    pf.space_after  = Pt(8)

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F6F8")
    pPr.append(shd)

    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "1")
        b.set(qn("w:color"), "B0B7C0")
        pBdr.append(b)
    pPr.append(pBdr)

    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x2F)


def add_image(doc, img_path, max_w=6.5):
    if not os.path.exists(img_path):
        para = doc.add_paragraph(f"[image missing: {img_path}]")
        add_bidi(para)
        for r in para.runs: r.italic = True
        return
    try:
        doc.add_picture(img_path, width=Inches(max_w))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as ex:
        para = doc.add_paragraph(f"[image error {img_path}: {ex}]")
        add_bidi(para)


def add_table_from_md(doc, header, rows):
    if not header: return
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3D6B")
        tcPr.append(shd)
        set_cell_rtl(cell)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            if c >= len(header): continue
            cell = table.rows[r].cells[c]
            cell.text = ""
            para = cell.paragraphs[0]
            add_runs_with_inline(para, val)
            set_cell_rtl(cell)


def add_runs_with_inline(paragraph, text):
    pos = 0
    pattern = re.compile(
        r"(\*\*(?P<bold>.+?)\*\*)"
        r"|(`(?P<code>[^`]+?)`)"
        r"|(?<!\*)\*(?P<italic>[^*\n]+?)\*(?!\*)"
    )
    while pos < len(text):
        m = pattern.search(text, pos)
        if not m:
            paragraph.add_run(text[pos:])
            return
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group("bold") is not None:
            r = paragraph.add_run(m.group("bold")); r.bold = True
        elif m.group("code") is not None:
            r = paragraph.add_run(m.group("code"))
            r.font.name = "Consolas"; r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif m.group("italic") is not None:
            r = paragraph.add_run(m.group("italic")); r.italic = True
        pos = m.end()


def process_markdown(md_text, doc):
    lines = md_text.splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]

        if line.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            add_code_block(doc, "\n".join(buf))
            continue

        if line.startswith("# "):     add_heading(doc, line[2:].strip(), 0); i += 1; continue
        if line.startswith("## "):    add_heading(doc, line[3:].strip(), 1); i += 1; continue
        if line.startswith("### "):   add_heading(doc, line[4:].strip(), 2); i += 1; continue
        if line.startswith("#### "):  add_heading(doc, line[5:].strip(), 3); i += 1; continue

        m_img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line)
        if m_img:
            alt, src = m_img.group(1), m_img.group(2)
            full = src if os.path.isabs(src) else os.path.join(ROOT, src.replace("/", os.sep))
            add_image(doc, full)
            if alt:
                cap = doc.add_paragraph(alt)
                add_bidi(cap)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cap.runs:
                    r.italic = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        if "|" in line and i + 1 < n and re.match(r"^\s*\|?\s*[:\-\s|]+\s*\|?\s*$", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            add_table_from_md(doc, header, rows)
            continue

        if re.match(r"^\s*[-*]\s+", line):
            while i < n and re.match(r"^\s*[-*]\s+", lines[i]):
                m = re.match(r"^\s*[-*]\s+(.+)$", lines[i])
                p = doc.add_paragraph(style="List Bullet")
                add_bidi(p)
                add_runs_with_inline(p, m.group(1))
                i += 1
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            while i < n and re.match(r"^\s*\d+\.\s+", lines[i]):
                m = re.match(r"^\s*\d+\.\s+(.+)$", lines[i])
                p = doc.add_paragraph(style="List Number")
                add_bidi(p)
                add_runs_with_inline(p, m.group(1))
                i += 1
            continue

        if line.startswith("> "):
            text = line[2:]
            p = doc.add_paragraph()
            add_bidi(p)
            pf = p.paragraph_format
            pf.left_indent = Cm(0.6)
            pf.right_indent = Cm(0.6)
            r = p.add_run(text)
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            right = OxmlElement("w:right")
            right.set(qn("w:val"), "single")
            right.set(qn("w:sz"), "12")
            right.set(qn("w:space"), "8")
            right.set(qn("w:color"), "F39C12")
            pBdr.append(right)
            pPr.append(pBdr)
            i += 1
            continue

        if re.match(r"^---+\s*$", line):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "888888")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        p = doc.add_paragraph()
        add_bidi(p)
        add_runs_with_inline(p, line)
        i += 1


def build_one(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        md = f.read()

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    pPr = style.element.get_or_add_pPr()
    pPr.append(OxmlElement("w:bidi"))

    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width  = Cm(21.0)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    process_markdown(md, doc)
    doc.save(docx_path)


def build():
    for md, dx in JOBS:
        if os.path.exists(md):
            build_one(md, dx)


if __name__ == "__main__":
    build()
