"""Build the Hebrew project book .docx from v6 markdown.
Force RTL at every level: section, paragraph, run, table."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"D:\yudb\ספר_פרויקט_B_Managed_FINAL_v6.md"
OUT = r"D:\yudb\ספר_פרויקט_B_Managed_FINAL_v6b.docx"

doc = Document()

# A4
section = doc.sections[0]
section.page_height = Inches(11.69)
section.page_width  = Inches(8.27)
section.left_margin = section.right_margin = Inches(0.8)
section.top_margin  = section.bottom_margin = Inches(0.8)

# Section-level RTL — flips entire document gutter to Hebrew
sectPr = section._sectPr
bidiSect = OxmlElement("w:bidi")
sectPr.append(bidiSect)
rtlGutter = OxmlElement("w:rtlGutter")
sectPr.append(rtlGutter)

# Default style — David CS for Hebrew, Calibri for Latin
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
rPr = style.element.get_or_add_rPr()
for tag in rPr.findall(qn("w:rFonts")):
    rPr.remove(tag)
rFonts = OxmlElement("w:rFonts")
rFonts.set(qn("w:cs"),    "David")
rFonts.set(qn("w:ascii"), "Calibri")
rFonts.set(qn("w:hAnsi"), "Calibri")
rFonts.set(qn("w:hint"),  "cs")
rPr.append(rFonts)
# Run-level RTL flag baked into Normal style
rtl = OxmlElement("w:rtl")
rPr.append(rtl)
lang = OxmlElement("w:lang")
lang.set(qn("w:val"),       "he-IL")
lang.set(qn("w:bidi"),      "he-IL")
rPr.append(lang)

# Same for every Heading style
for hname in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
    try:
        st = doc.styles[hname]
        rp = st.element.get_or_add_rPr()
        for tag in rp.findall(qn("w:rFonts")):
            rp.remove(tag)
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:cs"), "David")
        rf.set(qn("w:hint"), "cs")
        rp.append(rf)
        rp.append(OxmlElement("w:rtl"))
    except KeyError:
        pass

def set_p_rtl(p):
    """Paragraph-level: bidi + right-align + run RTL."""
    pPr = p._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        rPr = r._r.get_or_add_rPr()
        if rPr.find(qn("w:rtl")) is None:
            rPr.append(OxmlElement("w:rtl"))
        if rPr.find(qn("w:cs")) is None:
            rPr.append(OxmlElement("w:cs"))

def set_p_ltr(p):
    """For code blocks — keep LTR + monospace."""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    set_p_rtl(p)
    return p

def add_para(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    set_p_rtl(p)
    return p

def add_code(text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F2F2EE")
    pPr.append(shd)
    set_p_ltr(p)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x0B, 0x0B, 0x0F)
    return p

def add_table(rows):
    if not rows: return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # bidiVisual on the table flips column order to RTL
    tblPr = t._tbl.tblPr
    bidiV = OxmlElement("w:bidiVisual")
    tblPr.append(bidiV)
    for i, row in enumerate(rows):
        for j, c in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(c))
            if i == 0:
                r.bold = True
            set_p_rtl(p)
    return t

# ---------------- Parse markdown ----------------
with open(SRC, "r", encoding="utf-8") as f:
    md = f.read()

lines = md.split("\n")
i = 0
in_code = False
code_buf = []
table_buf = []
in_table = False

def flush_code():
    global code_buf
    if code_buf:
        add_code("\n".join(code_buf))
        code_buf = []

def flush_table():
    global table_buf, in_table
    parsed = []
    for row in table_buf:
        cells = [c.strip() for c in row.strip("|").split("|")]
        if all(re.match(r"^-+:?$|^:?-+:?$", c.replace(" ", "")) for c in cells if c):
            continue
        parsed.append(cells)
    if parsed:
        add_table(parsed)
    table_buf = []
    in_table = False

while i < len(lines):
    line = lines[i]

    if line.strip().startswith("```"):
        if not in_code:
            flush_table()
            in_code = True
        else:
            in_code = False
            flush_code()
        i += 1
        continue
    if in_code:
        code_buf.append(line)
        i += 1
        continue

    if line.lstrip().startswith("|") and "|" in line[2:]:
        flush_code()
        table_buf.append(line)
        in_table = True
        i += 1
        continue
    elif in_table:
        flush_table()

    m = re.match(r"^(#+)\s+(.+)$", line.strip())
    if m:
        flush_code(); flush_table()
        level = min(len(m.group(1)), 4)
        add_heading(m.group(2).strip(), level=level)
        i += 1
        continue

    if line.strip().startswith("---"):
        i += 1
        continue

    if line.strip().startswith("* ") or line.strip().startswith("- "):
        flush_code()
        text = re.sub(r"^[\s]*[\*\-]\s+", "", line)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text)
        set_p_rtl(p)
        i += 1
        continue

    if re.match(r"^\d+\.\s", line.strip()):
        flush_code()
        text = re.sub(r"^\d+\.\s+", "", line.strip())
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        p = doc.add_paragraph(style="List Number")
        p.add_run(text)
        set_p_rtl(p)
        i += 1
        continue

    if line.strip() == "":
        i += 1
        continue

    text = line.strip()
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    add_para(text)
    i += 1

flush_code()
flush_table()

doc.save(OUT)
print("done")
